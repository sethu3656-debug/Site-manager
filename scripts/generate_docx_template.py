import os
import docx

def create_template():
    doc = docx.Document()
    
    # Title
    doc.add_heading('Site Information Report', 0)
    
    # 1. Sites Section
    doc.add_paragraph('{% for site in sites %}')
    doc.add_heading('Site: {{ site.name }}', level=1)
    
    # Site Info Fields
    p = doc.add_paragraph()
    p.add_run('Description: ').bold = True
    p.add_run('{{ site.description }}\n')
    p.add_run('Latitude: ').bold = True
    p.add_run('{{ site.latitude }}  |  ')
    p.add_run('Longitude: ').bold = True
    p.add_run('{{ site.longitude }}\n')
    p.add_run('Height: ').bold = True
    p.add_run('{{ site.height }}\n')
    
    p.add_run('\nChain 1 Positions:\n').bold = True
    p.add_run('Position X: ').bold = True
    p.add_run('{{ site.x1 }}  |  ')
    p.add_run('Position Y: ').bold = True
    p.add_run('{{ site.y1 }}  |  ')
    p.add_run('Position Z: ').bold = True
    p.add_run('{{ site.z1 }}\n')
    
    p.add_run('\nChain 2 Positions:\n').bold = True
    p.add_run('Position X: ').bold = True
    p.add_run('{{ site.x2 }}  |  ')
    p.add_run('Position Y: ').bold = True
    p.add_run('{{ site.y2 }}  |  ')
    p.add_run('Position Z: ').bold = True
    p.add_run('{{ site.z2 }}\n')
    p.add_run('Latitude 2: ').bold = True
    p.add_run('{{ site.latitude2 }}  |  ')
    p.add_run('Longitude 2: ').bold = True
    p.add_run('{{ site.longitude2 }}  |  ')
    p.add_run('Height 2: ').bold = True
    p.add_run('{{ site.height2 }}\n')
    
    p.add_run('\nRF Delays:\n').bold = True
    p.add_run('L1 Delay: ').bold = True
    p.add_run('{{ site.l1_delay }} ns  |  ')
    p.add_run('S Delay: ').bold = True
    p.add_run('{{ site.s_delay }} ns  |  ')
    p.add_run('L5 Delay: ').bold = True
    p.add_run('{{ site.l5_delay }} ns\n')
    
    doc.add_paragraph('{% endfor %}')
    
    # 2. Racks Section
    doc.add_heading('Racks', level=1)
    # Check if racks are present in loop
    doc.add_paragraph('{% if racks %}')
    table_racks = doc.add_table(rows=2, cols=6)
    table_racks.style = 'Table Grid'
    hdr_cells = table_racks.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Serial'
    hdr_cells[2].text = 'Asset ID'
    hdr_cells[3].text = 'Site'
    hdr_cells[4].text = 'Height'
    hdr_cells[5].text = 'Devices'
    
    row_cells = table_racks.rows[1].cells
    row_cells[0].text = '{% for rack in racks %}{{ rack.name }}'
    row_cells[1].text = '{{ rack.serial }}'
    row_cells[2].text = '{{ rack.asset_id }}'
    row_cells[3].text = '{{ rack.site }}'
    row_cells[4].text = '{{ rack.height }}'
    row_cells[5].text = '{{ rack.devices }}{% endfor %}'
    doc.add_paragraph('{% endif %}')
    
    # 3. Equipment Section
    doc.add_heading('Equipment', level=1)
    doc.add_paragraph('{% if equipment %}')
    table_eq = doc.add_table(rows=2, cols=7)
    table_eq.style = 'Table Grid'
    hdr_eq = table_eq.rows[0].cells
    hdr_eq[0].text = 'Name'
    hdr_eq[1].text = 'Type'
    hdr_eq[2].text = 'Role'
    hdr_eq[3].text = 'Location'
    hdr_eq[4].text = 'Serial'
    hdr_eq[5].text = 'IP Address'
    hdr_eq[6].text = 'Status'
    
    row_eq = table_eq.rows[1].cells
    row_eq[0].text = '{% for eq in equipment %}{{ eq.name }}'
    row_eq[1].text = '{{ eq.device_type }}'
    row_eq[2].text = '{{ eq.role }}'
    row_eq[3].text = '{{ eq.location }}'
    row_eq[4].text = '{{ eq.serial }}'
    row_eq[5].text = '{{ eq.ip_address }}'
    row_eq[6].text = '{{ eq.status }}{% endfor %}'
    doc.add_paragraph('{% endif %}')
    
    # 4. Cables Section
    doc.add_heading('Cables', level=1)
    doc.add_paragraph('{% if cables %}')
    table_cables = doc.add_table(rows=2, cols=7)
    table_cables.style = 'Table Grid'
    hdr_cables = table_cables.rows[0].cells
    hdr_cables[0].text = 'Label'
    hdr_cables[1].text = 'Device A'
    hdr_cables[2].text = 'Port A'
    hdr_cables[3].text = 'Device B'
    hdr_cables[4].text = 'Port B'
    hdr_cables[5].text = 'Length'
    hdr_cables[6].text = 'Type'
    
    row_cables = table_cables.rows[1].cells
    row_cables[0].text = '{% for cable in cables %}{{ cable.label }}'
    row_cables[1].text = '{{ cable.dev_a }}'
    row_cables[2].text = '{{ cable.term_a }}'
    row_cables[3].text = '{{ cable.dev_b }}'
    row_cables[4].text = '{{ cable.term_b }}'
    row_cables[5].text = '{{ cable.length }}'
    row_cables[6].text = '{{ cable.type }}{% endfor %}'
    doc.add_paragraph('{% endif %}')
    
    # 5. IP Addresses Section
    doc.add_heading('IP Addresses', level=1)
    doc.add_paragraph('{% if ips %}')
    table_ips = doc.add_table(rows=2, cols=5)
    table_ips.style = 'Table Grid'
    hdr_ips = table_ips.rows[0].cells
    hdr_ips[0].text = 'S.No'
    hdr_ips[1].text = 'Device'
    hdr_ips[2].text = 'Interface'
    hdr_ips[3].text = 'IP Address'
    hdr_ips[4].text = 'Description'
    
    row_ips = table_ips.rows[1].cells
    row_ips[0].text = '{% for ip in ips %}{{ ip.s_no }}'
    row_ips[1].text = '{{ ip.device }}'
    row_ips[2].text = '{{ ip.interface }}'
    row_ips[3].text = '{{ ip.ip_address }}'
    row_ips[4].text = '{{ ip.description }}{% endfor %}'
    doc.add_paragraph('{% endif %}')
    
    # 6. Chains Section
    doc.add_heading('RF Delay Chains', level=1)
    doc.add_paragraph('{% if chains %}')
    doc.add_paragraph('{% for chain in chains %}')
    doc.add_heading('Chain: {{ chain.name }}', level=2)
    
    table_chains = doc.add_table(rows=2, cols=4)
    table_chains.style = 'Table Grid'
    hdr_chains = table_chains.rows[0].cells
    hdr_chains[0].text = 'Device'
    hdr_chains[1].text = 'L1 Delay (ns)'
    hdr_chains[2].text = 'L5 Delay (ns)'
    hdr_chains[3].text = 'S Delay (ns)'
    
    row_chains = table_chains.rows[1].cells
    row_chains[0].text = '{% for delay in chain.delays %}{{ delay.device }}'
    row_chains[1].text = '{{ delay.l1|format_num }}'
    row_chains[2].text = '{{ delay.l5|format_num }}'
    row_chains[3].text = '{{ delay.s|format_num }}{% endfor %}'
    
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('{% endif %}')
    
    # Create directory and save
    os.makedirs('netbox_site_docx', exist_ok=True)
    output_path = os.path.join('netbox_site_docx', 'Site Information.docx')
    doc.save(output_path)
    print(f"Jinja template successfully saved at: {output_path}")

if __name__ == '__main__':
    create_template()
